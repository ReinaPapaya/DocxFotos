import os
import io
import zipfile
from flask import Flask, request, render_template, send_file, jsonify
from werkzeug.utils import secure_filename
from PIL import Image
from docx import Document
from docx.shared import Inches, Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 100 * 1024 * 1024  # 100MB

# ---- Utilidades internas ----

def set_cell_background(cell, color_hex):
    """Pinta fondo de celda en tabla (para imágenes con fondo)"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex.replace('#', ''))
    tcPr.append(shd)

def add_caption(paragraph, text, numbering=False, num_prefix="Figura "):
    """Añade caption con numeración automática (estilo 'Caption' o normal)"""
    if numbering:
        # Contador global por sesión (se pasa como argumento)
        pass  # lo manejamos en la función principal
    run = paragraph.add_run(text)
    run.bold = True
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Intentar aplicar estilo Caption si existe
    try:
        paragraph.style = 'Caption'
    except:
        pass  # si no existe, queda normal

def insert_image_with_options(doc, img_stream, config, fig_num, caption_text):
    """
    Inserta una imagen en el documento según configuraciones.
    Retorna el párrafo o tabla donde se insertó (para posibles saltos).
    """
    # Abrir imagen con PIL para manejar transparencias
    pil_img = Image.open(img_stream)
    # Convertir a RGB si es necesario
    bg_color = config.get('fondo', 'transparente')
    if bg_color != 'transparente' and pil_img.mode in ('RGBA', 'LA', 'P'):
        # Crear fondo
        bg = Image.new('RGB', pil_img.size, bg_color)
        if pil_img.mode == 'P':
            pil_img = pil_img.convert('RGBA')
        bg.paste(pil_img, (0, 0), pil_img if pil_img.mode == 'RGBA' else None)
        pil_img = bg
    elif pil_img.mode != 'RGB':
        pil_img = pil_img.convert('RGB')
    
    # Guardar en bytes
    img_bytes = io.BytesIO()
    pil_img.save(img_bytes, format='PNG')
    img_bytes.seek(0)

    # Dimensiones
    width_val = config.get('ancho', 80)
    width_unit = config.get('unidad_ancho', '%')  # '%' o 'cm'
    if width_unit == '%':
        # Calcular ancho en cm asumiendo página A4 (21cm - márgenes)
        page_width_cm = 16.0  # aprox ancho útil
        width_inches = Inches((width_val / 100.0) * page_width_cm / 2.54)
    else:
        width_inches = Cm(float(width_val))

    keep_ratio = config.get('mantener_proporcion', True)
    
    # Crear párrafo o tabla
    alignment_map = {
        'izquierda': WD_ALIGN_PARAGRAPH.LEFT,
        'centrado': WD_ALIGN_PARAGRAPH.CENTER,
        'derecha': WD_ALIGN_PARAGRAPH.RIGHT
    }
    align = alignment_map.get(config.get('alineacion', 'centrado'), WD_ALIGN_PARAGRAPH.CENTER)
    
    # Si el pie es lateral: usar tabla de una fila y dos celdas
    pie_pos = config.get('pie_posicion', 'abajo')
    if pie_pos == 'lateral':
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        table.allow_autofit = False
        # Celda izquierda (imagen)
        cell_img = table.cell(0, 0)
        cell_img.width = int(width_inches) * 2  # ajuste
        # Celda derecha (caption)
        cell_cap = table.cell(0, 1)
        # Insertar imagen en celda izquierda
        paragraph_img = cell_img.paragraphs[0] if cell_img.paragraphs else cell_img.add_paragraph()
        run = paragraph_img.add_run()
        run.add_picture(img_bytes, width=width_inches)
        paragraph_img.alignment = align
        # Añadir caption en celda derecha
        p_cap = cell_cap.paragraphs[0] if cell_cap.paragraphs else cell_cap.add_paragraph()
        if config.get('numeracion', False):
            caption_full = f"Figura {fig_num}: {caption_text}"
        else:
            caption_full = caption_text
        add_caption(p_cap, caption_full, numbering=False)
        # Fondo de celda (para transparencias)
        if bg_color != 'transparente':
            set_cell_background(cell_img, bg_color)
        return table
    else:
        # Pie abajo: párrafo imagen + párrafo caption
        p_img = doc.add_paragraph()
        p_img.alignment = align
        run = p_img.add_run()
        run.add_picture(img_bytes, width=width_inches)
        # Si fondo no transparente, añadir sombreado al párrafo? no es trivial
        # Añadir caption
        p_cap = doc.add_paragraph()
        if config.get('numeracion', False):
            caption_full = f"Figura {fig_num}: {caption_text}"
        else:
            caption_full = caption_text
        add_caption(p_cap, caption_full, numbering=False)
        # Espaciado
        spacing = config.get('espaciado_pt', 12)
        p_cap.paragraph_format.space_before = Pt(spacing)
        p_cap.paragraph_format.space_after = Pt(spacing)
        return p_img

def find_and_replace_bookmark(doc, bookmark_text, replacement_elements):
    """
    Busca un párrafo que contenga exactamente bookmark_text (ej. '#INSERT#')
    y lo reemplaza con los elementos pasados (list de párrafos/tablas).
    """
    bookmark_text = bookmark_text.strip()
    for paragraph in doc.paragraphs:
        if bookmark_text in paragraph.text:
            # Insertar elementos justo antes de este párrafo
            parent = paragraph._element.getparent()
            for elem in replacement_elements:
                # elem es un objeto lxml (párrafo o tabla)
                parent.insert(parent.index(paragraph._element), elem)
            # Eliminar párrafo marcador
            parent.remove(paragraph._element)
            return True
    return False

# ---- Endpoints ----

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/generar', methods=['POST'])
def generar():
    # 1. Recibir archivos y configuración
    semilla = request.files.get('semilla')
    if not semilla:
        return jsonify({'error': 'Semilla requerida'}), 400
    imagenes = request.files.getlist('imagenes')
    config_json = request.form.get('config')
    import json
    config = json.loads(config_json) if config_json else {}

    # 2. Cargar documento semilla
    doc = Document(semilla)
    
    # 3. Preparar imágenes ordenadas
    # Las imágenes vienen en orden según frontend (drag&drop)
    # El frontend envía un campo 'orden' con lista de nombres o índices
    orden_str = request.form.get('orden', '')
    orden = json.loads(orden_str) if orden_str else []
    # Filtro por inclusión (checkbox)
    incluidos_str = request.form.get('incluidos', '')
    incluidos = json.loads(incluidos_str) if incluidos_str else []
    
    # Asociar cada imagen con su índice original
    img_map = {f'img_{i}': (i, imagenes[i]) for i in range(len(imagenes)) if imagenes[i].filename}
    
    # Reordenar según 'orden' (lista de keys tipo 'img_0')
    ordered_keys = orden if orden else list(img_map.keys())
    # Filtrar solo incluidos
    if incluidos:
        ordered_keys = [k for k in ordered_keys if k in incluidos]
    
    # 4. Insertar imágenes
    inserted_elements = []
    fig_counter = 1
    bookmark = config.get('bookmark', '#INSERT_IMAGES#')
    
    # Modo de ajuste: una por página o flujo
    one_per_page = config.get('modo_ajuste') == 'una_por_pagina'
    
    for key in ordered_keys:
        idx, file = img_map[key]
        # Obtener caption personalizado
        caption_type = config.get('caption_tipo', 'nombre_archivo')
        if caption_type == 'nombre_archivo':
            caption_text = file.filename.rsplit('.', 1)[0]
        elif caption_type == 'secuencial':
            caption_text = f"Imagen {idx+1}"
        else:  # 'personalizado'
            # El frontend envía captions personalizados en un dict
            captions_custom = json.loads(request.form.get('captions_custom', '{}'))
            caption_text = captions_custom.get(key, file.filename)
        
        # Insertar
        elem = insert_image_with_options(
            doc, file.stream, config,
            fig_counter if config.get('numeracion', False) else None,
            caption_text
        )
        inserted_elements.append(elem)
        
        # Salto de página si modo una por página (excepto último)
        if one_per_page and ordered_keys.index(key) < len(ordered_keys)-1:
            doc.add_page_break()
        
        fig_counter += 1
    
    # 5. Reemplazar bookmark
    # Los elementos insertados están en el documento al final, debemos moverlos al bookmark
    # Solución: recolectar los elementos recién insertados (están al final del doc)
    # En lugar de usar find_and_replace_bookmark, mejor insertar al final y luego mover no es trivial.
    # Simplificamos: buscamos el marcador y reemplazamos con los elementos recién creados.
    # Los elementos recién creados están en doc.paragraphs[-N:] y doc.tables[-M:]
    # Pero es más fiable: al insertar, guardamos referencias a los objetos lxml.
    # Como insert_image_with_options crea elementos y los añade al final, recuperamos los últimos.
    # Mejor enfoque: pasar doc y un "anchor" para insertar in situ.
    # Rediseño rápido: en lugar de bookmark, insertaremos al final (simplificación)
    # Pero el usuario pidió bookmark. Implementaremos una búsqueda y reemplazo con los elementos generados.
    
    # Como los elementos fueron añadidos al final, los extraemos.
    # Obtenemos el número de párrafos/tablas antes de insertar
    # Haremos una versión más robusta: recolectamos los elementos creados en una lista
    # Ya tenemos inserted_elements que son objetos lxml (internos de docx).
    # Pero no son lxml directamente, sino objetos de python-docx.
    # Para reemplazar, necesitamos acceder a ._element
    elem_lxml = [e._element for e in inserted_elements if hasattr(e, '_element')]
    # Buscar bookmark
    bookmark_found = False
    for p in doc.paragraphs:
        if bookmark in p.text:
            parent = p._element.getparent()
            for el in elem_lxml:
                parent.insert(parent.index(p._element), el)
            parent.remove(p._element)
            bookmark_found = True
            break
    if not bookmark_found:
        # Si no se encuentra, añadir al final
        for el in elem_lxml:
            doc.element.body.append(el)
    
    # 6. Guardar en memoria
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    
    # 7. Devolver
    nombre_salida = request.form.get('nombre_salida', 'documento_generado.docx')
    if not nombre_salida.endswith('.docx'):
        nombre_salida += '.docx'
    return send_file(
        output,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=nombre_salida
    )

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)